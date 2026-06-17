"""DOCX export (python-docx).

Render a ReportModel as a downloadable Word document mirroring the snapshot: branded header, the
full 2-column table (13 MVP + 12 metric rows via presentation.all_rows), and a clickable Medicare
Care Compare hyperlink.
"""
from __future__ import annotations

from io import BytesIO

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Pt, RGBColor

from medelite import config, presentation
from medelite.models import ReportModel

_HYPERLINK_REL = "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink"
_BRAND_RGB = RGBColor.from_string(config.BRAND_COLOR.lstrip("#"))
_LINK_RGB = "1A73E8"
_LABEL_FILL = "F5F6FA"


def _shade(cell, fill: str) -> None:
    """Apply a background fill to a table cell (python-docx has no direct API for this)."""
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:fill"), fill)
    cell._tc.get_or_add_tcPr().append(shd)


def _add_hyperlink(paragraph, url: str, text: str) -> None:
    """Append a real clickable hyperlink run to a paragraph (standard python-docx OOXML recipe)."""
    r_id = paragraph.part.relate_to(url, _HYPERLINK_REL, is_external=True)
    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(qn("r:id"), r_id)
    run = OxmlElement("w:r")
    rpr = OxmlElement("w:rPr")
    color = OxmlElement("w:color")
    color.set(qn("w:val"), _LINK_RGB)
    rpr.append(color)
    underline = OxmlElement("w:u")
    underline.set(qn("w:val"), "single")
    rpr.append(underline)
    run.append(rpr)
    text_el = OxmlElement("w:t")
    text_el.text = text
    run.append(text_el)
    hyperlink.append(run)
    paragraph._p.append(hyperlink)


def build_docx(rep: ReportModel) -> bytes:
    """Render the report to DOCX bytes (clickable Medicare hyperlink)."""
    doc = Document()

    head = doc.add_paragraph()
    head.alignment = WD_ALIGN_PARAGRAPH.CENTER
    brand_run = head.add_run(config.BRAND_PLATFORM)
    brand_run.bold = True
    brand_run.font.size = Pt(20)
    brand_run.font.color.rgb = _BRAND_RGB
    tag_run = head.add_run("  " + config.BRAND_TAGLINE)
    tag_run.font.size = Pt(10)
    tag_run.font.color.rgb = RGBColor.from_string("666666")

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_run = title.add_run(config.REPORT_TITLE)
    title_run.bold = True
    title_run.font.size = Pt(12)

    name = doc.add_paragraph()
    name_run = name.add_run(rep.facility_name)
    name_run.bold = True
    name_run.font.size = Pt(13)

    rows = presentation.all_rows(rep)
    table = doc.add_table(rows=len(rows), cols=2)
    table.style = "Table Grid"
    for i, (label, value) in enumerate(rows):
        lcell, vcell = table.rows[i].cells
        lrun = lcell.paragraphs[0].add_run(label)
        lrun.bold = True
        _shade(lcell, _LABEL_FILL)
        vcell.paragraphs[0].add_run(value)

    doc.add_paragraph()
    _add_hyperlink(doc.add_paragraph(), rep.medicare_url, "View on Medicare Care Compare")

    if not rep.cms_record_found:
        note = doc.add_paragraph()
        note_run = note.add_run("No CMS record found for this CCN; values shown are from manual inputs.")
        note_run.font.size = Pt(8)
        note_run.font.color.rgb = RGBColor.from_string("999999")

    buf = BytesIO()
    doc.save(buf)
    return buf.getvalue()
